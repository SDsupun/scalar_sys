import base64
import os
import subprocess

from docx import Document
from docx.shared import RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from mysql import connector

#os.environ["SCALAR_WD"] = "/home/svradmin/Documents/scalar_sys"
script_dir = os.environ["SCALAR_WD"]


class ReportFormat1:

    def __init__(self):

        get_ticket_format = '''SELECT * FROM scalar_system.ticket_format WHERE format='Report' '''
        get_company_info = '''SELECT company_name, address, telephone, email, fax FROM scalar_system.sys_setting'''
        db_conn = connector.connect(host="localhost", user="scal_user", password="tH@r@236", database="scalar_system")

        # need to get from the database
        cursor = db_conn.cursor()
        cursor.execute(get_ticket_format)
        ticket_format = cursor.fetchone()

        self.document = Document()
        self.section = self.document.sections[0]
        
        new_width, new_height = self.section.page_height, self.section.page_width
        self.section.orientation = WD_ORIENT.LANDSCAPE
        self.section.page_height = new_height
        self.section.page_width = new_width

        self.section.left_margin = Inches(ticket_format[5])
        self.section.right_margin = Inches(ticket_format[7])
        self.section.top_margin = Inches(ticket_format[6])
        self.section.bottom_margin = Inches(ticket_format[8])

        cursor = db_conn.cursor()
        cursor.execute(get_company_info)
        company_info_base64 = cursor.fetchone()

        self.company = base64.b64decode(company_info_base64[0].encode('ascii')).decode('ascii')
        self.company_address = base64.b64decode(company_info_base64[1].encode('ascii')).decode('ascii')
        self.company_tel = base64.b64decode(company_info_base64[2].encode('ascii')).decode('ascii')
        self.company_email = base64.b64decode(company_info_base64[3].encode('ascii')).decode('ascii')
        self.company_fax = base64.b64decode(company_info_base64[4].encode('ascii')).decode('ascii')

        self.old_ticket_present = False

        # set doc style
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Roman'
        self.font.size = Pt(8)

        db_conn.disconnect()
        return

    def set_search_info(self, search_info, records, column_count):

        if len(records) <= 1:
            return False

        search_header = f"{self.company} \n {self.company_address} \n {self.company_tel} | {self.company_email} | " \
                        f"{self.company_fax}"

        heading = self.document.add_heading(search_header, level=2)
        heading.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph(search_info)

        table = self.document.add_table(rows=0, cols=column_count)
        table.style = 'TableGrid'

        for r_row in records:
            row_cells = table.add_row().cells
            for c in range(column_count):
                row_cells[c].text = r_row[c]

        return True

    def print_record(self):
        self.document.save(f'{script_dir}/ticket/record.docx')
        try:
            # /../ticket/new_ticket.docx for linux
            print_job = subprocess.call(['libreoffice', '-p', f'{script_dir}/ticket/record.docx'])
            if print_job == 0:
                return True
            else:
                return False
        except FileNotFoundError:
            return False
        except PermissionError:
            return False
