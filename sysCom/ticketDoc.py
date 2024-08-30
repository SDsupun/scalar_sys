import base64
import os
import subprocess

from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches, Pt
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from mysql import connector

#os.environ["SCALAR_WD"] = "/home/svradmin/Documents/scalar_sys"
script_dir = os.environ["SCALAR_WD"]


class TicketFormat1:

    def __init__(self):

        self.document = Document()
        self.section = self.document.sections[0]

        get_ticket_format = '''SELECT * FROM scalar_system.ticket_format WHERE format='Format-1' '''
        get_company_info = '''SELECT company_name, address, telephone, email, fax FROM scalar_system.sys_setting'''
        db_conn = connector.connect(host="localhost", user="scal_user", password="tH@r@236", database="scalar_system")

        # need to get from the database
        cursor = db_conn.cursor()
        cursor.execute(get_ticket_format)
        ticket_format = cursor.fetchone()

        self.section.page_height = Inches(ticket_format[3])
        self.section.page_width = Inches(ticket_format[4])

        self.section.left_margin = Inches(ticket_format[5])
        self.section.right_margin = Inches(ticket_format[7])
        self.section.top_margin = Inches(ticket_format[6])
        self.section.bottom_margin = Inches(ticket_format[8])
        self.custom_info = base64.b64decode(ticket_format[1].encode('ascii')).decode('ascii')
        self.footer = base64.b64decode(ticket_format[2].encode('ascii')).decode('ascii')

        cursor = db_conn.cursor()
        cursor.execute(get_company_info)
        company_info_base64 = cursor.fetchone()

        self.company = base64.b64decode(company_info_base64[0].encode('ascii')).decode('ascii')
        self.company_address = base64.b64decode(company_info_base64[1].encode('ascii')).decode('ascii')
        self.company_tel = base64.b64decode(company_info_base64[2].encode('ascii')).decode('ascii')
        self.company_email = base64.b64decode(company_info_base64[3].encode('ascii')).decode('ascii')
        self.company_fax = base64.b64decode(company_info_base64[4].encode('ascii')).decode('ascii')

        self.old_ticket_present = False

        # set doc style
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Roman'
        self.font.size = Pt(10)

        db_conn.disconnect()

        return

    def set_ticket_info(self, ticket_detail):
        ticket_header = f"{self.company} \n {self.company_address} \n {self.company_tel} | {self.company_email} | " \
                        f"{self.company_fax}"

        ticket_info = f"{self.custom_info}"

        ticket_footer = f"\n{self.footer}"

        records = (
            ('DATE & TIME', f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 'BILL NO', f'{ticket_detail[0]}'),
            ('VEHICLE NO', f'{ticket_detail[1]}', '', ''),
            ('', '1st WEIGHT', f'{ticket_detail[2]}', ''),
            ('', '2nd WEIGHT', f'{ticket_detail[3]}', ''),
            ('INCOMING TIME', f'{ticket_detail[4]}', '', ''),
            ('OUTGOING TIME', f'{ticket_detail[5]}', '', ''),
            ('', 'NET WEIGHT', f'{ticket_detail[6]}', '')
        )

        heading = self.document.add_heading(ticket_header, level=2)
        heading.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = self.document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col1, col2, col3, col4 in records:
            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3
            row_cells[3].text = col4
            if 'weight' in col1.lower() + col2.lower() + col3.lower() + col4.lower():
                run1 = row_cells[0].paragraphs[0].runs
                run1[0].font.bold = True
                run1[0].font.size = Pt(12)
                run2 = row_cells[1].paragraphs[0].runs
                run2[0].font.bold = True
                run2[0].font.size = Pt(12)
                run3 = row_cells[2].paragraphs[0].runs
                run3[0].font.bold = True
                run3[0].font.size = Pt(12)

        self.document.add_paragraph(ticket_info)
        self.document.add_paragraph(ticket_footer)

    def print_docx(self, ticket_no, weight_no):
        self.old_ticket_present = True
        try:
            os.remove(f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no - 1}.docx")
        except FileNotFoundError:
            pass
        self.document.save(f'{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx')
        try:
            print_job = subprocess.call(['libreoffice', '-p',
                                         f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx"])
            if print_job == 0:
                return True
            else:
                return False
        except FileNotFoundError:
            return False
        except PermissionError:
            return False

    def print_previous_docx(self, ticket_no, weight_no):
        if self.old_ticket_present:
            try:
                print_job = subprocess.call(['libreoffice', '-p',
                                             f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no-1}.docx"])
                if print_job == 0:
                    return True
                else:
                    return False
            except FileNotFoundError:
                return False
            except PermissionError:
                return False


class TicketFormat2:

    def __init__(self):

        self.document = Document()
        self.section = self.document.sections[0]

        get_ticket_format = '''SELECT * FROM scalar_system.ticket_format WHERE format='Format-2' '''
        get_company_info = '''SELECT company_name, address, telephone, email, fax FROM scalar_system.sys_setting'''
        db_conn = connector.connect(host="localhost", user="scal_user", password="tH@r@236", database="scalar_system")

        # need to get from the database
        cursor = db_conn.cursor()
        cursor.execute(get_ticket_format)
        ticket_format = cursor.fetchone()

        self.section.page_height = Inches(ticket_format[3])
        self.section.page_width = Inches(ticket_format[4])

        self.section.left_margin = Inches(ticket_format[5])
        self.section.right_margin = Inches(ticket_format[7])
        self.section.top_margin = Inches(ticket_format[6])
        self.section.bottom_margin = Inches(ticket_format[8])
        self.custom_info = base64.b64decode(ticket_format[1].encode('ascii')).decode('ascii')
        self.footer = base64.b64decode(ticket_format[2].encode('ascii')).decode('ascii')

        cursor = db_conn.cursor()
        cursor.execute(get_company_info)
        company_info_base64 = cursor.fetchone()

        self.company = base64.b64decode(company_info_base64[0].encode('ascii')).decode('ascii')
        self.company_address = base64.b64decode(company_info_base64[1].encode('ascii')).decode('ascii')
        self.company_tel = base64.b64decode(company_info_base64[2].encode('ascii')).decode('ascii')
        self.company_email = base64.b64decode(company_info_base64[3].encode('ascii')).decode('ascii')
        self.company_fax = base64.b64decode(company_info_base64[4].encode('ascii')).decode('ascii')

        self.old_ticket_present = False

        # set doc style
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Roman'
        self.font.size = Pt(10)

        db_conn.disconnect()
        return

    def set_ticket_info(self, ticket_detail):
        ticket_header = f"{self.company} \n {self.company_address} \n {self.company_tel} | {self.company_email} " \
                        f"| {self.company_fax}"

        ticket_info = f"{self.custom_info}"

        ticket_footer = f"{self.footer}"

        records1 = (
            ('CUSTOMER NAME', f'{ticket_detail[7]}'),
            ('PRODUCT', f'{ticket_detail[8]}')
        )

        records2 = (
            ('', '1st WEIGHT', '2nd WEIGHT'),
            ('VEHICLE', f'{ticket_detail[1]}', f'{ticket_detail[1]}'),
            ('DATE TIME', f'{ticket_detail[4]}', f'{ticket_detail[5]}'),
            ('CONSEC NUM', f'{ticket_detail[0]}', f'{ticket_detail[0]}'),
            ('WEIGHT', f'{ticket_detail[2]}', f'{ticket_detail[3]}'),
            ('NET WEIGHT', f'{ticket_detail[6]}', '')
        )

        heading = self.document.add_heading(ticket_header, level=2)
        heading.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph('')

        table = self.document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col1, col2 in records1:
            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2

        table2 = self.document.add_table(rows=0, cols=3)

        for col1, col2, col3 in records2:
            row_cells = table2.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3
            if 'weight' in col1.lower() + col2.lower() + col3.lower():
                run1 = row_cells[0].paragraphs[0].runs
                run1[0].font.bold = True
                run1[0].font.size = Pt(12)
                run2 = row_cells[1].paragraphs[0].runs
                run2[0].font.bold = True
                run2[0].font.size = Pt(12)
                run3 = row_cells[2].paragraphs[0].runs
                run3[0].font.bold = True
                run3[0].font.size = Pt(12)

        self.document.add_paragraph('')
        self.document.add_paragraph(ticket_info)
        self.document.add_paragraph(ticket_footer)

    def print_docx(self, ticket_no, weight_no):
        self.old_ticket_present = True
        try:
            os.remove(f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no - 1}.docx")
        except FileNotFoundError:
            pass
        self.document.save(f'{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx')
        try:
            print_job = subprocess.call(['libreoffice', '-p',
                                         f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx"])
            if print_job == 0:
                return True
            else:
                return False
        except FileNotFoundError:
            return False
        except PermissionError:
            return False

    def print_previous_docx(self, ticket_no, weight_no):
        if self.old_ticket_present:
            try:
                print_job = subprocess.call(['libreoffice', '-p',
                                             f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no-1}.docx"])
                if print_job == 0:
                    return True
                else:
                    return False
            except FileNotFoundError:
                return False
            except PermissionError:
                return False


class TicketFormat3:

    def __init__(self):

        self.document = Document()
        self.section = self.document.sections[0]

        get_ticket_format = '''SELECT * FROM scalar_system.ticket_format WHERE format='Format-3' '''
        get_company_info = '''SELECT company_name, address, telephone, email, fax FROM scalar_system.sys_setting'''
        db_conn = connector.connect(host="localhost", user="scal_user", password="tH@r@236", database="scalar_system")

        # need to get from the database
        cursor = db_conn.cursor()
        cursor.execute(get_ticket_format)
        ticket_format = cursor.fetchone()

        self.section.page_height = Inches(ticket_format[3])
        self.section.page_width = Inches(ticket_format[4])

        self.section.left_margin = Inches(ticket_format[5])
        self.section.right_margin = Inches(ticket_format[7])
        self.section.top_margin = Inches(ticket_format[6])
        self.section.bottom_margin = Inches(ticket_format[8])
        self.custom_info = base64.b64decode(ticket_format[1].encode('ascii')).decode('ascii')
        self.footer = base64.b64decode(ticket_format[2].encode('ascii')).decode('ascii')

        cursor = db_conn.cursor()
        cursor.execute(get_company_info)
        company_info_base64 = cursor.fetchone()

        self.company = base64.b64decode(company_info_base64[0].encode('ascii')).decode('ascii')
        self.company_address = base64.b64decode(company_info_base64[1].encode('ascii')).decode('ascii')
        self.company_tel = base64.b64decode(company_info_base64[2].encode('ascii')).decode('ascii')
        self.company_email = base64.b64decode(company_info_base64[3].encode('ascii')).decode('ascii')
        self.company_fax = base64.b64decode(company_info_base64[4].encode('ascii')).decode('ascii')

        self.old_ticket_present = False

        # set doc style
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Roman'
        self.font.size = Pt(10)

        db_conn.disconnect()
        return

    def set_ticket_info(self, records):
        ticket_header = f"{self.company} \n {self.company_address} \n {self.company_tel} | {self.company_email} " \
                        f"| {self.company_fax}"

        ticket_info = f"{self.custom_info}"

        ticket_footer = f"\n{self.footer}"

        heading = self.document.add_heading(ticket_header, level=2)
        heading.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        table = self.document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col1, col2 in records:
            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            if 'weight' in col1.lower() + col2.lower():
                run1 = row_cells[0].paragraphs[0].runs
                run1[0].font.bold = True
                run1[0].font.size = Pt(12)
                run2 = row_cells[1].paragraphs[0].runs
                run2[0].font.bold = True
                run2[0].font.size = Pt(12)

        self.document.add_paragraph(ticket_info)
        self.document.add_paragraph(ticket_footer)

    def print_docx(self, ticket_no, weight_no):
        self.old_ticket_present = True
        try:
            os.remove(f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no - 1}.docx")
        except FileNotFoundError:
            pass
        self.document.save(f'{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx')
        try:
            print_job = subprocess.call(['libreoffice', '-p',
                                         f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx"])
            if print_job == 0:
                return True
            else:
                return False
        except FileNotFoundError:
            return False
        except PermissionError:
            return False

    def print_previous_docx(self, ticket_no, weight_no):
        if self.old_ticket_present:
            try:
                print_job = subprocess.call(['libreoffice', '-p',
                                             f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no-1}.docx"])
                if print_job == 0:
                    return True
                else:
                    return False
            except FileNotFoundError:
                return False
            except PermissionError:
                return False


class TicketFormat4:

    def __init__(self):

        self.document = Document()
        self.section = self.document.sections[0]

        get_ticket_format = '''SELECT * FROM scalar_system.ticket_format WHERE format='Format-4' '''
        get_company_info = '''SELECT company_name, address, telephone, email, fax FROM scalar_system.sys_setting'''
        db_conn = connector.connect(host="localhost", user="scal_user", password="tH@r@236", database="scalar_system")

        # need to get from the database
        cursor = db_conn.cursor()
        cursor.execute(get_ticket_format)
        ticket_format = cursor.fetchone()

        self.section.page_height = Inches(ticket_format[3])
        self.section.page_width = Inches(ticket_format[4])

        self.section.left_margin = Inches(ticket_format[5])
        self.section.right_margin = Inches(ticket_format[7])
        self.section.top_margin = Inches(ticket_format[6])
        self.section.bottom_margin = Inches(ticket_format[8])
        self.custom_info = base64.b64decode(ticket_format[1].encode('ascii')).decode('ascii')
        self.footer = base64.b64decode(ticket_format[2].encode('ascii')).decode('ascii')

        cursor = db_conn.cursor()
        cursor.execute(get_company_info)
        company_info_base64 = cursor.fetchone()

        self.company = base64.b64decode(company_info_base64[0].encode('ascii')).decode('ascii')
        self.company_address = base64.b64decode(company_info_base64[1].encode('ascii')).decode('ascii')
        self.company_tel = base64.b64decode(company_info_base64[2].encode('ascii')).decode('ascii')
        self.company_email = base64.b64decode(company_info_base64[3].encode('ascii')).decode('ascii')
        self.company_fax = base64.b64decode(company_info_base64[4].encode('ascii')).decode('ascii')

        self.old_ticket_present = False

        # set doc style
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Roman'
        self.font.size = Pt(10)

        db_conn.disconnect()
        return

    def set_ticket_info(self, ticket_detail):
        ticket_header = f"{self.company} \n {self.company_address} \n {self.company_tel} | {self.company_email} " \
                        f"| {self.company_fax}"

        ticket_info = f"{self.custom_info}"

        ticket_footer = f"{self.footer}"

        records1 = (
            ('TICKET NUMBER', f'{ticket_detail[0]}', 'MATERIAL', f'{ticket_detail[8]}'),
            ('VEHICLE NO', f'{ticket_detail[1]}', 'CUSTOMER', f'{ticket_detail[7]}'),
            ('DATE-TIME', f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', '', '')
        )

        records2 = (
            ('FIRST WEIGHT', f'{ticket_detail[2]}', 'IN TIME', f'{ticket_detail[4]}'),
            ('SECOND WEIGHT', f'{ticket_detail[3]}', 'OUT TIME', f'{ticket_detail[5]}'),
            ('NET WEIGHT', f'{ticket_detail[6]}', '', '')
        )

        heading = self.document.add_heading(ticket_header, level=2)
        heading.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph('')

        table = self.document.add_table(rows=0, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col1, col2, col3, col4 in records1:
            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3
            row_cells[3].text = col4

        table2 = self.document.add_table(rows=0, cols=4)

        for col1, col2, col3, col4 in records2:
            row_cells = table2.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3
            row_cells[3].text = col4

            run2 = row_cells[1].paragraphs[0].runs
            run2[0].font.bold = True
            run2[0].font.size = Pt(12)

        self.document.add_paragraph('')
        self.document.add_paragraph(ticket_info)
        self.document.add_paragraph(ticket_footer)

    def print_docx(self, ticket_no, weight_no):
        self.old_ticket_present = True
        try:
            os.remove(f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no - 1}.docx")
        except FileNotFoundError:
            pass
        self.document.save(f'{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx')
        try:
            print_job = subprocess.call(['libreoffice', '-p',
                                         f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no}.docx"])
            if print_job == 0:
                return True
            else:
                return False
        except FileNotFoundError:
            return False
        except PermissionError:
            return False

    def print_previous_docx(self, ticket_no, weight_no):
        if self.old_ticket_present:
            try:
                print_job = subprocess.call(['libreoffice', '-p',
                                             f"{script_dir}/ticket/ticket-{weight_no}-{ticket_no-1}.docx"])
                if print_job == 0:
                    return True
                else:
                    return False
            except FileNotFoundError:
                return False
            except PermissionError:
                return False
